"""Unit tests for TailoredResumeService — fakes ResumeExportService's
gather_resume_data, private storage, and the LLM; no database, no real
document rendering dependencies beyond the real (pure, no-I/O)
build_resume_docx/build_resume_pdf functions.
"""

from __future__ import annotations

import json
import uuid
from datetime import UTC, date, datetime

import pytest

from app.application.jd_tailoring.tailored_resume_service import (
    TailoredResumeService,
    _parse_tailored_content,
    _repair_raw_control_chars_in_json_strings,
)
from app.core.exceptions import CareerCompassError, NotFoundError
from app.adapters.documents.resume_data import ResumeData
from app.domain.career_profile.entities import CareerProfile, Experience
from app.domain.identity.entities import User
from app.domain.jd_tailoring.entities import JdTailoringSession

pytestmark = pytest.mark.unit


class FakeJdTailoringSessionRepository:
    def __init__(self, session: JdTailoringSession) -> None:
        self.sessions = {session.id: session}
        self.updated: list[JdTailoringSession] = []

    async def get_by_id(self, tenant_id: uuid.UUID, session_id: uuid.UUID) -> JdTailoringSession | None:
        s = self.sessions.get(session_id)
        return s if s and s.tenant_id == tenant_id else None

    async def update(self, session: JdTailoringSession) -> JdTailoringSession:
        self.sessions[session.id] = session
        self.updated.append(session)
        return session


class FakeResumeExportService:
    def __init__(self, data: ResumeData) -> None:
        self._data = data

    async def gather_resume_data(self, *, tenant_id, user_id, target_role_id):
        return self._data.profile, self._data

    async def gather_resume_data_with_master_fallback(self, *, tenant_id, user_id, target_role_id):
        return self._data.profile, self._data


class FakeStorage:
    def __init__(self) -> None:
        self.uploaded: dict[str, bytes] = {}

    async def upload_private(self, *, key: str, content: bytes, content_type: str) -> None:
        self.uploaded[key] = content

    async def get_presigned_url(self, *, key: str, expires_in_seconds: int = 300, download_filename=None) -> str:
        return f"https://storage.example.com/{key}?signed=1"

    async def delete_private(self, *, key: str) -> None:
        self.uploaded.pop(key, None)


class FakeLLMService:
    def __init__(self, response_text: str | None = None, fail: bool = False) -> None:
        self._response_text = response_text
        self._fail = fail
        self.last_input_variables: dict[str, str] | None = None

    async def generate(self, *, use_case: str, input_variables: dict[str, str], **kwargs) -> str:
        self.last_input_variables = input_variables
        if self._fail:
            raise CareerCompassError("simulated provider failure")
        return self._response_text or "{}"


def _make_resume_data(tenant_id: uuid.UUID, user_id: uuid.UUID) -> tuple[ResumeData, uuid.UUID]:
    now = datetime.now(UTC)
    profile_id = uuid.uuid4()
    profile = CareerProfile(
        id=profile_id,
        tenant_id=tenant_id,
        user_id=user_id,
        current_version=1,
        headline="Software Engineer",
        summary="An experienced engineer.",
        career_readiness_score=None,
        photo_url=None,
        core_competencies=[],
        created_at=now,
        updated_at=now,
    )
    user = User(
        id=user_id,
        tenant_id=tenant_id,
        org_id=None,
        email="jordan@example.com",
        salutation=None,
        first_name="Jordan",
        last_name="Rivera",
        hashed_password="x",
        status="active",
        mfa_enabled=False,
        created_at=now,
        updated_at=now,
    )
    experience_id = uuid.uuid4()
    experience = Experience(
        id=experience_id,
        tenant_id=tenant_id,
        career_profile_id=profile_id,
        title="Engineer",
        company="Acme",
        location=None,
        start_date=date(2020, 1, 1),
        end_date=None,
        description="Built things.",
        display_order=1,
        created_at=now,
        updated_at=now,
    )
    data = ResumeData(
        profile=profile,
        user=user,
        experiences=[experience],
        educations=[],
        certifications=[],
        career_highlights=[],
        key_achievements=[],
        career_goals=[],
        recommendations=[],
        role_label=None,
    )
    return data, experience_id


def _make_session(tenant_id: uuid.UUID, user_id: uuid.UUID, **overrides) -> JdTailoringSession:
    now = datetime.now(UTC)
    defaults: dict[str, object] = dict(
        id=uuid.uuid4(),
        tenant_id=tenant_id,
        user_id=user_id,
        source_type="custom",
        jd_text="A job description.",
        created_at=now,
        updated_at=now,
    )
    defaults.update(overrides)
    return JdTailoringSession(**defaults)  # type: ignore[arg-type]


class TestGenerate:
    async def test_successful_generation_stores_a_docx_and_returns_a_url(self) -> None:
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, experience_id = _make_resume_data(tenant_id, user_id)
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        storage = FakeStorage()
        llm = FakeLLMService(
            response_text=json.dumps(
                {
                    "headline": "Senior Software Engineer",
                    "summary": "A tailored summary.",
                    "experience_bullets": [
                        {"id": str(experience_id), "bullets": ["Did X", "Did Y"]}
                    ],
                }
            )
        )
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), storage, llm
        )

        updated, url = await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="docx"
        )

        assert updated.tailored_resume_status == "generated"
        assert updated.tailored_resume_error is None
        assert updated.tailored_resume_docx_key is not None
        assert updated.tailored_resume_docx_key.startswith(f"tailored-resumes/{tenant_id}/{session.id}/")
        assert url is not None
        assert updated.tailored_resume_docx_key in storage.uploaded

    async def test_failed_generation_does_not_erase_a_previously_good_key(self) -> None:
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, _ = _make_resume_data(tenant_id, user_id)
        session = _make_session(
            tenant_id,
            user_id,
            tailored_resume_docx_key="tailored-resumes/t/s/resume.docx",
            tailored_resume_status="generated",
        )
        sessions_repo = FakeJdTailoringSessionRepository(session)
        storage = FakeStorage()
        llm = FakeLLMService(fail=True)
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), storage, llm
        )

        updated, url = await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="docx"
        )

        assert updated.tailored_resume_status == "failed"
        assert updated.tailored_resume_error is not None
        # The previously-good key survives a failed regenerate.
        assert updated.tailored_resume_docx_key == "tailored-resumes/t/s/resume.docx"
        # get_download_urls-style behavior: a URL is still returned for
        # the existing key, since generate() re-derives the URL from
        # whatever key ends up persisted.
        assert url is not None

    async def test_malformed_json_is_treated_as_a_failure_not_a_crash(self) -> None:
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, _ = _make_resume_data(tenant_id, user_id)
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        storage = FakeStorage()
        llm = FakeLLMService(response_text="not valid json")
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), storage, llm
        )

        updated, url = await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="pdf"
        )

        assert updated.tailored_resume_status == "failed"
        assert url is None

    async def test_a_raw_newline_inside_a_bullet_string_no_longer_fails_generation(
        self,
    ) -> None:
        # Confirmed live (2026-08-19): a real "Regenerate Word" attempt
        # failed with "Expecting ',' delimiter" — a classic symptom of
        # an LLM emitting a real, un-escaped control character inside a
        # JSON string value. This is the exact failure shape the new
        # repair pass (_repair_raw_control_chars_in_json_strings) exists
        # to recover from.
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, experience_id = _make_resume_data(tenant_id, user_id)
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        storage = FakeStorage()
        broken_json = (
            '{"headline": "Engineer", "summary": "S",'
            f' "experience_bullets": [{{"id": "{experience_id}",'
            ' "bullets": ["A bullet with a\nliteral newline inside it"]}]}'
        )
        llm = FakeLLMService(response_text=broken_json)
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), storage, llm
        )

        updated, url = await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="docx"
        )

        assert updated.tailored_resume_status == "generated"
        assert updated.tailored_resume_error is None
        assert url is not None

    async def test_cannot_generate_for_another_users_session(self) -> None:
        tenant_id = uuid.uuid4()
        owner, other = uuid.uuid4(), uuid.uuid4()
        data, _ = _make_resume_data(tenant_id, owner)
        session = _make_session(tenant_id, owner)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), FakeStorage(), FakeLLMService()
        )

        with pytest.raises(NotFoundError):
            await service.generate(
                tenant_id=tenant_id, user_id=other, session_id=session.id, format="docx"
            )


class TestCapsExperiencesSentToTheModel:
    async def test_only_the_first_n_experiences_are_sent_to_the_llm(self) -> None:
        # Confirmed live (2026-08-19): once gather_resume_data_with_master_fallback
        # started feeding a real profile's *entire* experience list into
        # this prompt, a heavily-used real account (many experiences)
        # blew the fixed max_tokens output budget mid-JSON. Uncapped
        # input was the root cause, not too-small a budget.
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, _ = _make_resume_data(tenant_id, user_id)
        # _make_resume_data seeds exactly one experience — add many more
        # so the cap is actually exercised.
        now = datetime.now(UTC)
        extra = [
            Experience(
                id=uuid.uuid4(),
                tenant_id=tenant_id,
                career_profile_id=data.profile.id,
                title=f"Role {i}",
                company=f"Company {i}",
                location=None,
                start_date=date(2015, 1, 1),
                end_date=None,
                description="Did things.",
                display_order=i,
                created_at=now,
                updated_at=now,
            )
            for i in range(15)
        ]
        data = ResumeData(
            profile=data.profile,
            user=data.user,
            experiences=data.experiences + extra,
            educations=[],
            certifications=[],
            career_highlights=[],
            key_achievements=[],
            career_goals=[],
            recommendations=[],
            role_label=None,
        )
        assert len(data.experiences) == 16  # well beyond the cap
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        llm = FakeLLMService(
            response_text=json.dumps(
                {"headline": "H", "summary": "S", "experience_bullets": []}
            )
        )
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), FakeStorage(), llm
        )

        await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="docx"
        )

        assert llm.last_input_variables is not None
        sent = json.loads(llm.last_input_variables["experiences_json"])
        assert len(sent) == 8

    async def test_a_long_description_is_truncated_and_flattened_to_plain_text(self) -> None:
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, experience_id = _make_resume_data(tenant_id, user_id)
        long_html = "<ul><li>" + ("Delivered impact. " * 100) + "</li></ul>"
        data.experiences[0].description = long_html
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        llm = FakeLLMService(
            response_text=json.dumps(
                {"headline": "H", "summary": "S", "experience_bullets": []}
            )
        )
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), FakeStorage(), llm
        )

        await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="docx"
        )

        assert llm.last_input_variables is not None
        sent = json.loads(llm.last_input_variables["experiences_json"])
        assert len(sent) == 1
        assert len(sent[0]["description"]) <= 500
        assert "<" not in sent[0]["description"]


class TestNormalizesUnsafeHyphens:
    async def test_replaces_non_breaking_hyphen_with_a_plain_ascii_hyphen(self) -> None:
        # Confirmed live (2026-08-19): CP1252/WinAnsiEncoding — what
        # reportlab's base-14 Helvetica font uses for the PDF export —
        # has no code point for U+2011 (non-breaking hyphen), silently
        # corrupting it into an unrelated glyph in the generated PDF.
        # Em-dash/en-dash are unaffected (real CP1252 code points), so
        # this only needs to catch the hyphen-family characters.
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, experience_id = _make_resume_data(tenant_id, user_id)
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        storage = FakeStorage()
        llm = FakeLLMService(
            response_text=json.dumps(
                {
                    "headline": "High‑Throughput Systems Engineer",
                    "summary": "Built job‑scheduling systems — real em-dash stays untouched.",
                    "experience_bullets": [
                        {"id": str(experience_id), "bullets": ["Owned end‑to‑end delivery"]}
                    ],
                }
            )
        )
        service = TailoredResumeService(sessions_repo, FakeResumeExportService(data), storage, llm)

        updated, _url = await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="docx"
        )

        # tailored_resume_content is deliberately the RAW, un-normalized
        # AI output (audit trail) — the normalization only applies to
        # what actually gets rendered into the document, so assert
        # against the real generated DOCX bytes, not that field.
        import io

        from docx import Document

        assert updated.tailored_resume_docx_key is not None
        docx_bytes = storage.uploaded[updated.tailored_resume_docx_key]
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "‑" not in full_text
        assert "High-Throughput" in full_text
        assert "—" in full_text  # real em-dash left untouched
        assert "Owned end-to-end delivery" in full_text


class TestNormalizesModelBulletPrefixes:
    async def test_strips_a_bullet_prefix_the_model_already_added_before_adding_its_own(
        self,
    ) -> None:
        # Confirmed live (2026-08-19): despite the prompt only asking for
        # the "• " convention, a real generation's bullets already came
        # back prefixed with "• " — this code unconditionally adds its
        # own "• " on top, so without stripping first the result is a
        # double-prefixed "• • text", and plain_text_to_rich_html only
        # strips the outer layer, leaving a stray literal "•" inside the
        # rendered list item's own text.
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, experience_id = _make_resume_data(tenant_id, user_id)
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        storage = FakeStorage()
        llm = FakeLLMService(
            response_text=json.dumps(
                {
                    "headline": "Engineer",
                    "summary": "Summary.",
                    "experience_bullets": [
                        {
                            "id": str(experience_id),
                            "bullets": [
                                "• Completed a core task delivering expected outcomes",
                                "- Executed an additional task",
                            ],
                        }
                    ],
                }
            )
        )
        service = TailoredResumeService(sessions_repo, FakeResumeExportService(data), storage, llm)

        updated, _url = await service.generate(
            tenant_id=tenant_id, user_id=user_id, session_id=session.id, format="docx"
        )

        import io

        from docx import Document

        assert updated.tailored_resume_docx_key is not None
        docx_bytes = storage.uploaded[updated.tailored_resume_docx_key]
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Exactly the bullet text, no residual "•"/"-" prefix character
        # left over from the model's own attempt at the convention.
        assert "Completed a core task delivering expected outcomes" in full_text
        assert "• •" not in full_text
        assert "• -" not in full_text


class TestRepairsRawControlCharsInJsonStrings:
    def test_escapes_a_raw_newline_inside_a_string_value(self) -> None:
        broken = '{"headline": "Line one\nLine two", "summary": "S"}'
        repaired = _repair_raw_control_chars_in_json_strings(broken)
        parsed = json.loads(repaired)
        assert parsed["headline"] == "Line one\nLine two"

    def test_escapes_a_raw_tab_and_carriage_return_inside_a_string(self) -> None:
        broken = '{"a": "col1\tcol2", "b": "line1\rline2"}'
        parsed = json.loads(_repair_raw_control_chars_in_json_strings(broken))
        assert parsed == {"a": "col1\tcol2", "b": "line1\rline2"}

    def test_leaves_already_valid_json_semantically_unchanged(self) -> None:
        valid = json.dumps({"headline": "H", "summary": "S", "experience_bullets": []})
        repaired = _repair_raw_control_chars_in_json_strings(valid)
        assert json.loads(repaired) == json.loads(valid)

    def test_does_not_touch_whitespace_between_tokens_outside_strings(self) -> None:
        spaced = '{\n  "headline": "H",\n  "summary": "S"\n}'
        repaired = _repair_raw_control_chars_in_json_strings(spaced)
        assert json.loads(repaired) == {"headline": "H", "summary": "S"}

    def test_respects_backslash_escaped_quotes_when_tracking_string_state(self) -> None:
        # A \" inside a string must not be mistaken for the string's
        # closing quote — if it were, the newline right after it would
        # be treated as "outside a string" and left unescaped, still
        # producing invalid JSON.
        broken = '{"headline": "She said \\"hi\\"\nthen left"}'
        repaired = _repair_raw_control_chars_in_json_strings(broken)
        parsed = json.loads(repaired)
        assert parsed["headline"] == 'She said "hi"\nthen left'


class TestParseTailoredContentRecovery:
    def test_recovers_from_a_raw_newline_inside_a_bullet(self) -> None:
        broken = (
            '{"headline": "H", "summary": "S", "experience_bullets": '
            '[{"id": "abc", "bullets": ["Line one\nstill one bullet"]}]}'
        )
        parsed = _parse_tailored_content(broken)
        assert parsed["headline"] == "H"

    def test_still_raises_for_genuinely_unparseable_text(self) -> None:
        with pytest.raises(ValueError):
            _parse_tailored_content("not json at all")


class TestGetDownloadUrls:
    async def test_returns_none_for_formats_never_generated(self) -> None:
        tenant_id, user_id = uuid.uuid4(), uuid.uuid4()
        data, _ = _make_resume_data(tenant_id, user_id)
        session = _make_session(tenant_id, user_id)
        sessions_repo = FakeJdTailoringSessionRepository(session)
        service = TailoredResumeService(
            sessions_repo, FakeResumeExportService(data), FakeStorage(), FakeLLMService()
        )

        urls = await service.get_download_urls(session)

        assert urls.docx_url is None
        assert urls.pdf_url is None
