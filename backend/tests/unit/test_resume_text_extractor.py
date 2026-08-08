"""Unit tests for PdfDocxTextExtractor.

Builds real, valid DOCX files in-memory with python-docx's own writer
API — the same library used to read them — so these tests exercise the
real adapter against a real file format, not a hand-rolled stand-in.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from app.adapters.parsing.resume_text_extractor import DOCX_CONTENT_TYPE, PdfDocxTextExtractor
from app.core.exceptions import CareerCompassError


def _build_docx(builder) -> bytes:
    document = Document()
    builder(document)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.unit
class TestExtractDocx:
    def test_plain_paragraphs_are_extracted(self) -> None:
        def builder(doc):
            doc.add_paragraph("Jordan Rivera")
            doc.add_paragraph("Senior Engineer")

        content = _build_docx(builder)
        extractor = PdfDocxTextExtractor()

        text = extractor.extract_text(content=content, content_type=DOCX_CONTENT_TYPE)

        assert "Jordan Rivera" in text
        assert "Senior Engineer" in text

    def test_table_content_is_extracted_not_silently_dropped(self) -> None:
        """Real bug caught live: the original extractor only read
        document.paragraphs, which never included table content at all.
        Word resume templates commonly use tables for layout (e.g. a
        title/company/date row per job) — that content was completely
        invisible to the LLM, not a prompt-quality issue.
        """

        def builder(doc):
            doc.add_paragraph("PROFESSIONAL EXPERIENCE")
            table = doc.add_table(rows=1, cols=3)
            cells = table.rows[0].cells
            cells[0].text = "Acme Corp"
            cells[1].text = "Staff Engineer"
            cells[2].text = "Jan 2025 - Present"

        content = _build_docx(builder)
        extractor = PdfDocxTextExtractor()

        text = extractor.extract_text(content=content, content_type=DOCX_CONTENT_TYPE)

        assert "Acme Corp" in text
        assert "Staff Engineer" in text
        assert "Jan 2025 - Present" in text

    def test_paragraphs_and_tables_are_extracted_in_document_order(self) -> None:
        def builder(doc):
            doc.add_paragraph("EDUCATION")
            table = doc.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = "State University"
            doc.add_paragraph("CERTIFICATIONS")

        content = _build_docx(builder)
        extractor = PdfDocxTextExtractor()

        text = extractor.extract_text(content=content, content_type=DOCX_CONTENT_TYPE)
        lines = [line for line in text.split("\n") if line.strip()]

        assert lines == ["EDUCATION", "State University", "CERTIFICATIONS"]

    def test_empty_cells_and_paragraphs_are_skipped(self) -> None:
        def builder(doc):
            doc.add_paragraph("")
            doc.add_paragraph("Real content")
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Filled"
            table.rows[0].cells[1].text = ""

        content = _build_docx(builder)
        extractor = PdfDocxTextExtractor()

        text = extractor.extract_text(content=content, content_type=DOCX_CONTENT_TYPE)
        lines = [line for line in text.split("\n") if line.strip()]

        assert lines == ["Real content", "Filled"]

    def test_malformed_docx_raises_text_extraction_error(self) -> None:
        extractor = PdfDocxTextExtractor()

        with pytest.raises(CareerCompassError) as exc_info:
            extractor.extract_text(content=b"not a real docx file", content_type=DOCX_CONTENT_TYPE)

        assert exc_info.value.code == "RESUME_TEXT_EXTRACTION_FAILED"
