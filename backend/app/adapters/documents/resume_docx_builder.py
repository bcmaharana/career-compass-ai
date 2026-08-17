"""Generates a formatted .docx resume from a ResumeData bundle.

python-docx is already a project dependency (used today for *reading*
uploaded resumes in resume_text_extractor.py) — its Document API is
bidirectional, so no new dependency is needed to *write* one.

Section order follows the specific profile's own
resolve_resume_section_order(profile.section_order) — the same user
-customizable order already shown on the Career Profile page for
*that* profile (Master or a Target Role Profile each order
independently, since section_order lives on the CareerProfile row
itself) — not a hardcoded sequence. Requested live after an earlier
version always rendered a fixed order regardless of how the person had
actually arranged their profile.
"""

from __future__ import annotations

import io
from collections.abc import Callable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor

from app.adapters.documents.resume_data import (
    ResumeData,
    contact_line_parts,
    education_degree_line,
    format_date_range,
    group_competencies_by_category,
    resolve_resume_section_order,
)
from app.adapters.documents.rich_text_export import parse_rich_text, plain_text


def _add_rich_blocks(
    document: DocumentObject,
    html: str | None,
    *,
    exclude: tuple[str | None, ...] = (),
    default_style: str | None = None,
) -> None:
    """Renders sanitized rich-text HTML (see rich_text_export.py) as one
    docx paragraph per block — bullet blocks always use "List Bullet"
    regardless of `default_style`, which only applies to plain blocks
    (e.g. "Intense Quote" for Recommendations). Bold/italic/underline/
    color are applied per run within each paragraph, so mixed formatting
    within one line survives. `exclude` drops any block whose flattened
    text is an exact (case/whitespace-insensitive) duplicate of one of
    the given reference strings — see resume_data.py's description_lines
    docstring for why (e.g. Education's own constructed degree line).
    """
    excluded = {ref.strip().casefold() for ref in exclude if ref}
    for block in parse_rich_text(html):
        block_text = "".join(run.text for run in block.runs).strip()
        if block_text.casefold() in excluded:
            continue
        style = "List Bullet" if block.bullet else default_style
        paragraph = document.add_paragraph(style=style)
        if block.indent:
            paragraph.paragraph_format.left_indent = Inches(0.4)
        for run in block.runs:
            docx_run = paragraph.add_run(run.text)
            docx_run.bold = run.bold
            docx_run.italic = run.italic
            docx_run.underline = run.underline
            if run.color:
                docx_run.font.color.rgb = RGBColor.from_string(run.color.lstrip("#"))


def _render_core_competencies(document: DocumentObject, data: ResumeData) -> None:
    if not data.profile.core_competencies:
        return
    document.add_heading("Core Competencies", level=1)
    for category, items in group_competencies_by_category(data.profile.core_competencies):
        group_p = document.add_paragraph()
        group_p.add_run(f"{category or 'Uncategorized'}: ").bold = True
        group_p.add_run(", ".join(item.name for item in items))


def _render_experience(document: DocumentObject, data: ResumeData) -> None:
    if not data.experiences:
        return
    document.add_heading("Professional Experience", level=1)
    for exp in data.experiences:
        title_p = document.add_paragraph()
        title_p.add_run(exp.title).bold = True
        title_p.add_run(f" — {exp.company}")
        if exp.location:
            title_p.add_run(f" ({exp.location})")
        date_p = document.add_paragraph(format_date_range(exp.start_date, exp.end_date))
        date_p.runs[0].italic = True
        _add_rich_blocks(document, exp.description)


def _render_education(document: DocumentObject, data: ResumeData) -> None:
    if not data.educations:
        return
    document.add_heading("Education", level=1)
    for edu in data.educations:
        degree_line = education_degree_line(edu)
        title_p = document.add_paragraph()
        title_p.add_run(edu.institution).bold = True
        if degree_line:
            title_p.add_run(f" — {degree_line}")
        if edu.start_date or edu.end_date:
            date_p = document.add_paragraph(format_date_range(edu.start_date, edu.end_date))
            date_p.runs[0].italic = True
        _add_rich_blocks(document, edu.description, exclude=(degree_line,))


def _render_certifications(document: DocumentObject, data: ResumeData) -> None:
    if not data.certifications:
        return
    document.add_heading("Certifications", level=1)
    for cert in data.certifications:
        p = document.add_paragraph()
        p.add_run(cert.name).bold = True
        p.add_run(f" — {cert.issuing_organization}")
        if cert.issue_date:
            p.add_run(f" ({cert.issue_date.strftime('%b %Y')})")


def _render_career_highlights(document: DocumentObject, data: ResumeData) -> None:
    if not data.career_highlights:
        return
    document.add_heading("Career Highlights", level=1)
    for highlight in data.career_highlights:
        document.add_paragraph(highlight.title, style="List Bullet")


def _render_key_achievements(document: DocumentObject, data: ResumeData) -> None:
    if not data.key_achievements:
        return
    document.add_heading("Key Achievements", level=1)
    for achievement in data.key_achievements:
        p = document.add_paragraph(style="List Bullet")
        if achievement.company:
            p.add_run(f"{achievement.company} — ").bold = True
        p.add_run(achievement.title).bold = True
        _add_rich_blocks(document, achievement.description)


def _render_career_goals(document: DocumentObject, data: ResumeData) -> None:
    if not data.career_goals:
        return
    document.add_heading("Career Goals", level=1)
    for goal in data.career_goals:
        p = document.add_paragraph()
        p.add_run(goal.target_role).bold = True
        if goal.target_date:
            p.add_run(f" (Target: {goal.target_date.strftime('%b %Y')})")
        _add_rich_blocks(document, goal.description)


def _render_recommendations(document: DocumentObject, data: ResumeData) -> None:
    if not data.recommendations:
        return
    document.add_heading("Recommendations", level=1)
    for endorsement in data.recommendations:
        _add_rich_blocks(document, endorsement.content, default_style="Intense Quote")
        attribution_p = document.add_paragraph()
        attribution_run = attribution_p.add_run(f"— {endorsement.recommender_name}")
        attribution_run.bold = True
        detail_parts = [
            part for part in (endorsement.recommender_title, endorsement.relationship) if part
        ]
        if detail_parts:
            attribution_p.add_run(f", {', '.join(detail_parts)}")


_SECTION_RENDERERS: dict[str, Callable[[DocumentObject, ResumeData], None]] = {
    "core_competencies": _render_core_competencies,
    "experience": _render_experience,
    "education": _render_education,
    "certifications": _render_certifications,
    "career_highlights": _render_career_highlights,
    "key_achievements": _render_key_achievements,
    "career_goals": _render_career_goals,
    "recommendations": _render_recommendations,
}


def build_resume_docx(data: ResumeData) -> bytes:
    document = Document()

    name = document.add_heading(data.user.display_name or data.user.email, level=0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_parts = [part for part in (data.role_label, plain_text(data.profile.headline)) if part]
    if subtitle_parts:
        subtitle = document.add_paragraph(" — ".join(subtitle_parts))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True

    contact_parts = contact_line_parts(data.user)
    if contact_parts:
        contact = document.add_paragraph(" | ".join(contact_parts))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if data.profile.summary:
        document.add_heading("Summary", level=1)
        _add_rich_blocks(document, data.profile.summary)

    for section_key in resolve_resume_section_order(data.profile.section_order):
        _SECTION_RENDERERS[section_key](document, data)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
