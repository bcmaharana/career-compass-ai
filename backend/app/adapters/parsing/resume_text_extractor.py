"""PDF/DOCX text extraction adapter.

Implements ResumeTextExtractor (app/domain/resume_intelligence/text_extraction.py).
Both pdfplumber and python-docx are synchronous, CPU-bound libraries —
callers must run extract_text() via asyncio.to_thread, the same
sync-wrapped-in-thread convention S3ObjectStorageRepository uses for
boto3.
"""

from __future__ import annotations

from collections.abc import Iterator
from io import BytesIO

import pdfplumber
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.exceptions import CareerCompassError

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class TextExtractionError(CareerCompassError):
    code = "RESUME_TEXT_EXTRACTION_FAILED"


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    """A paragraph is a genuine bulleted/numbered list item only if it
    carries Word's own list-numbering properties (`w:numPr`) — checked
    directly rather than via the paragraph's style *name*, since
    real-world resumes commonly apply a generic "List Paragraph" style
    to bullet items rather than a distinctly-named "List Bullet" one
    (verified live against a real resume: its actual bullet lines all
    reported `style.name == "List Paragraph"`, identical to what a
    plain intro sentence could also use — `numPr` is the one signal
    Word itself relies on to actually render the bullet/number glyph,
    regardless of style name).
    """
    p_pr = paragraph._p.pPr  # noqa: SLF001 — python-docx has no public numbering accessor
    return p_pr is not None and p_pr.find(qn("w:numPr")) is not None


def _iter_block_items(document: DocumentObject) -> Iterator[Paragraph | Table]:
    """Yields paragraphs and tables in actual document order.

    python-docx's top-level `document.paragraphs` and `document.tables`
    properties are two separate flat lists that each skip the other
    element type entirely — `.paragraphs` never included table content
    at all (a real bug: Word resume templates very commonly use tables
    for layout — a job-title/date column pair, a skills grid — and that
    content was silently invisible to the LLM, not a prompt/model
    quality issue at all). Walking the raw body XML and dispatching on
    each child's tag is the standard way to recover both, in their
    original relative order.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


class PdfDocxTextExtractor:
    def extract_text(self, *, content: bytes, content_type: str) -> str:
        if content_type == PDF_CONTENT_TYPE:
            return self._extract_pdf(content)
        if content_type == DOCX_CONTENT_TYPE:
            return self._extract_docx(content)
        raise TextExtractionError(f"Unsupported resume content type: {content_type!r}.")

    def _extract_pdf(self, content: bytes) -> str:
        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
        except Exception as exc:  # pdfplumber/pdfminer raise a variety of untyped errors
            raise TextExtractionError(f"Failed to parse PDF: {exc}") from exc
        return "\n\n".join(pages).strip()

    def _extract_docx(self, content: bytes) -> str:
        try:
            document = Document(BytesIO(content))
        # PackageNotFoundError covers "valid zip, wrong internal
        # structure"; a file that isn't a zip at all (garbage bytes, a
        # renamed non-.docx file) raises zipfile.BadZipFile instead,
        # which docx.opc's own reader doesn't translate — caught live by
        # this adapter's own test suite, not by review. python-docx has
        # no single documented exception type for "not a valid docx" the
        # way pdfplumber effectively doesn't either, so this matches
        # _extract_pdf's broad Exception catch for the same reason.
        except Exception as exc:
            raise TextExtractionError(f"Failed to parse DOCX: {exc}") from exc

        parts: list[str] = []
        for block in _iter_block_items(document):
            if isinstance(block, Paragraph):
                if block.text.strip():
                    # "• " marks a paragraph that Word itself rendered as
                    # a real bullet/numbered list item — preserved as a
                    # literal prefix so this distinction survives into the
                    # LLM's input (a plain intro sentence ahead of a
                    # role's bullets, and the bullets themselves, would
                    # otherwise be indistinguishable free text by the time
                    # they reach the prompt).
                    prefix = "• " if _is_list_paragraph(block) else ""
                    parts.append(prefix + block.text)
            else:
                for row in block.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
